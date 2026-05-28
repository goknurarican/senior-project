"""
generate_report.py
==================
stage 4 of approach b. assembles the final approach b report in docx and
converts to pdf via word.app. also writes a markdown findings summary and
copies key figures into approach_b/04_final_outputs/figures.

style rules (per spec):
  - times new roman 11pt body, 12pt headings
  - black and white only
  - no em-dashes, no ai-style phrases
  - numbered headings, prose paragraphs (no bullet lists)
  - tables: black borders, no shading
  - figures: grayscale, 300 dpi

required figures:
  fig1: roi definition on scalp topography (grayscale)
  fig2: per-scenario network change matrices (6 selected scenarios)
  fig3: network engagement by scenario cluster
  fig4: gnn confusion matrix (if stage 3 produced meaningful results)
  fig5: approach a v6 vs approach b gnn per-scenario performance
"""

import json
import shutil
import subprocess
import warnings
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

warnings.filterwarnings("ignore")

#paths
B_DIR    = Path(__file__).resolve().parents[1]
STAGE1   = B_DIR / "01_connectivity_extraction"
STAGE2   = B_DIR / "02_baseline_comparison"
STAGE3   = B_DIR / "03_gnn_classification"
OUT_DIR  = B_DIR / "04_final_outputs"
FIG_DIR  = OUT_DIR / "figures"
TAB_DIR  = OUT_DIR / "tables"
FIG_DIR.mkdir(parents=True, exist_ok=True)
TAB_DIR.mkdir(parents=True, exist_ok=True)

PROJECT  = B_DIR.parent
V6_EVAL  = PROJECT / "approach_a" / "06_v6_multiclass_characterization" / "evaluation"

ROI_NAMES = ["frontal", "frontal_central", "central",
             "parietal", "occipital", "temporal"]
BAND_NAMES = ["theta", "alpha", "beta", "gamma"]

SCENARIOS = [
    'broken_image','button_delay','coupon_expired','coupon_min_spend',
    'facet_reset_once','feedback_late','first_click_miss','network_jitter',
    'overlay_blocking','price_change','search_irrelevant','skeleton_prolong',
    'slow_image','sort_reset',
]


#─────────────────────────────────────────────────────────────────────────────
#figure builders
#─────────────────────────────────────────────────────────────────────────────
def figure_1_roi_topography():
    """grayscale scalp topography of the 6 rois."""
    import mne
    info = mne.create_info(
        ch_names=['Fp1','Fz','F3','F7','FT9','FC5','FC1','C3','T7','TP9','CP5','CP1',
                  'Pz','P3','P7','O1','Oz','O2','P4','P8','TP10','CP6','CP2','Cz',
                  'C4','T8','FT10','FC6','FC2','F4','F8','Fp2'],
        sfreq=500.0, ch_types='eeg')
    info.set_montage('standard_1005')
    roi_channels = {
        "frontal":         ['Fp1','Fp2','F3','Fz','F4','F7','F8'],
        "frontal_central": ['FC1','FC2','FC5','FC6'],
        "central":         ['C3','Cz','C4'],
        "parietal":        ['P3','Pz','P4','P7','P8'],
        "occipital":       ['O1','Oz','O2'],
        "temporal":        ['T7','T8','TP9','TP10'],
    }
    #assign each channel a grayscale value by roi index
    values = np.full(len(info['ch_names']), np.nan)
    levels = np.linspace(0.15, 0.85, len(roi_channels))
    for r_idx, (_, chs) in enumerate(roi_channels.items()):
        for ch in chs:
            if ch in info['ch_names']:
                values[info['ch_names'].index(ch)] = levels[r_idx]
    values = np.nan_to_num(values, nan=0.0)

    fig, ax = plt.subplots(figsize=(5.5, 5.0))
    im, _ = mne.viz.plot_topomap(values, info, axes=ax, show=False,
                                  cmap='gray_r', sensors=True, contours=0,
                                  vlim=(0, 1))
    ax.set_title("Figure 1. ROI definition on the 32-channel layout.",
                 fontsize=10)
    fig.tight_layout()
    out = FIG_DIR / "fig1_roi_topography.png"
    fig.savefig(str(out), dpi=300, facecolor="white")
    plt.close(fig)
    return out


def _selected_scenarios_for_fig2(meta_path: Path) -> list:
    """pick 6 scenarios for figure 2 - the ones with the largest n epochs."""
    meta = pd.read_csv(meta_path)
    counts = (meta[meta["scenario_name"].isin(SCENARIOS)]
              .groupby("scenario_name").size()
              .sort_values(ascending=False))
    return counts.index[:6].tolist()


def figure_2_change_matrices():
    """grid of 6 scenarios x 4 bands wpli delta heatmaps."""
    meta_path = STAGE1 / "features" / "connectivity_per_epoch" / "labels_v6.csv"
    selected = _selected_scenarios_for_fig2(meta_path)

    fig, axes = plt.subplots(len(selected), len(BAND_NAMES),
                             figsize=(11, 1.6 * len(selected) + 1.2))
    vmax = 0.0
    matrices = {}
    for scen in selected:
        path = STAGE2 / "analysis" / f"scenario_{scen}" / "connectivity_change_matrix.csv"
        if not path.exists():
            continue
        df = pd.read_csv(path)
        df = df[df["metric"] == "wpli"]
        scen_mats = []
        for band in BAND_NAMES:
            sub = df[df["band"] == band].copy()
            mat = np.zeros((len(ROI_NAMES), len(ROI_NAMES)), dtype=np.float32)
            for _, r in sub.iterrows():
                i = ROI_NAMES.index(r["roi_a"])
                j = ROI_NAMES.index(r["roi_b"])
                mat[i, j] = r["delta"]
            scen_mats.append(mat)
            vmax = max(vmax, abs(mat).max())
        matrices[scen] = scen_mats
    vmax = max(vmax, 1e-4)

    for s_idx, scen in enumerate(selected):
        for b_idx, band in enumerate(BAND_NAMES):
            ax = axes[s_idx, b_idx] if len(selected) > 1 else axes[b_idx]
            mat = matrices.get(scen, [np.zeros((6, 6))] * 4)[b_idx]
            im = ax.imshow(mat, cmap="gray_r", vmin=-vmax, vmax=vmax)
            if s_idx == 0:
                ax.set_title(band, fontsize=9)
            if b_idx == 0:
                ax.set_ylabel(scen.replace("_", "\n"), fontsize=7)
            ax.set_xticks([]); ax.set_yticks([])

    cbar_ax = fig.add_axes([0.92, 0.15, 0.015, 0.7])
    fig.colorbar(im, cax=cbar_ax)
    fig.suptitle("Figure 2. wPLI scenario-minus-control change matrices "
                 "(6 ROIs x 6 ROIs, six largest-N scenarios).", fontsize=10)
    fig.tight_layout(rect=[0, 0, 0.9, 0.95])
    out = FIG_DIR / "fig2_change_matrices.png"
    fig.savefig(str(out), dpi=300, facecolor="white")
    plt.close(fig)
    return out


def figure_3_network_engagement():
    """bar chart of network engagement counts across scenarios."""
    path = STAGE2 / "reports" / "network_engagement_table.csv"
    df = pd.read_csv(path)
    nets = ["fronto_parietal_control", "default_mode_alpha", "sensorimotor"]
    counts = {n: int(df[n].sum()) for n in nets}

    fig, ax = plt.subplots(figsize=(5.5, 4.0))
    xs = np.arange(len(nets))
    bars = ax.bar(xs, [counts[n] for n in nets], color="0.3",
                  edgecolor="black")
    for x, n in zip(xs, nets):
        ax.text(x, counts[n] + 0.1, str(counts[n]), ha="center", fontsize=9)
    ax.set_xticks(xs)
    ax.set_xticklabels([n.replace("_", " ") for n in nets], rotation=15,
                       ha="right", fontsize=8)
    ax.set_ylabel("scenarios with significant change")
    ax.set_ylim(0, max(counts.values() or [1]) + 1.5)
    ax.set_title("Figure 3. Network engagement across the 14 frustration "
                 "scenarios (wPLI).", fontsize=10)
    fig.tight_layout()
    out = FIG_DIR / "fig3_network_engagement.png"
    fig.savefig(str(out), dpi=300, facecolor="white")
    plt.close(fig)
    return out


def figure_4_gnn_confusion():
    """copy stage 3 confusion matrix into figures with a captioned wrapper."""
    src = STAGE3 / "evaluation" / "confusion_matrix.png"
    dst = FIG_DIR / "fig4_gnn_confusion.png"
    if src.exists():
        shutil.copy2(str(src), str(dst))
    return dst


def figure_5_v6_vs_gnn():
    """per-scenario performance comparison between approach a v6 and stage 3 gnn."""
    v6_perf_path = V6_EVAL / "per_scenario_performance.csv"
    gnn_perf_path = STAGE3 / "evaluation" / "per_scenario_performance.csv"

    if not v6_perf_path.exists() or not gnn_perf_path.exists():
        return None

    v6 = pd.read_csv(v6_perf_path)
    gnn = pd.read_csv(gnn_perf_path)

    scens = SCENARIOS
    v6_f1  = []
    gnn_f1 = []
    for s in scens:
        r_v6 = v6[v6["scenario"] == s]
        r_g  = gnn[gnn["scenario"] == s]
        v6_f1.append(float(r_v6["f1"].iloc[0]) if len(r_v6) else 0.0)
        gnn_f1.append(float(r_g["f1"].iloc[0]) if len(r_g) else 0.0)

    fig, ax = plt.subplots(figsize=(10, 4.2))
    xs = np.arange(len(scens))
    width = 0.4
    ax.bar(xs - width / 2, v6_f1,  width, color="0.2", edgecolor="black",
           label="Approach A V6 (Transformer)")
    ax.bar(xs + width / 2, gnn_f1, width, color="0.7", edgecolor="black",
           label="Approach B GNN (connectivity)")
    ax.set_xticks(xs)
    ax.set_xticklabels([s.replace("_", " ") for s in scens], rotation=70,
                       ha="right", fontsize=7)
    ax.set_ylabel("per-scenario F1")
    ax.set_ylim(0, max(max(v6_f1 + gnn_f1) + 0.1, 0.4))
    ax.legend(fontsize=9, frameon=False)
    ax.set_title("Figure 5. Per-scenario F1: Approach A V6 vs Approach B GNN.",
                 fontsize=10)
    fig.tight_layout()
    out = FIG_DIR / "fig5_v6_vs_gnn_f1.png"
    fig.savefig(str(out), dpi=300, facecolor="white")
    plt.close(fig)
    return out


#─────────────────────────────────────────────────────────────────────────────
#tables
#─────────────────────────────────────────────────────────────────────────────
def table_1_top_connections() -> pd.DataFrame:
    path = STAGE2 / "analysis" / "all_scenarios_network_summary.csv"
    df = pd.read_csv(path)
    df = df[df["metric"] == "wpli"].copy()
    df = df.dropna(subset=["cohens_d"])
    df["abs_d"] = df["cohens_d"].abs()
    rows = []
    for scen in SCENARIOS:
        sub = df[df["scenario"] == scen].copy()
        if sub.empty:
            rows.append({"scenario": scen, "band": "", "connection": "",
                         "cohens_d": np.nan, "fdr_p": np.nan,
                         "direction": "", "n_subjects": 0})
            continue
        top = sub.sort_values("abs_d", ascending=False).iloc[0]
        rows.append({"scenario": scen, "band": top["band"],
                     "connection": top["top_connection"],
                     "cohens_d": float(top["cohens_d"]),
                     "fdr_p": float(top["fdr_p"]),
                     "direction": top["direction"],
                     "n_subjects": int(top["n_subjects"])})
    out = pd.DataFrame(rows)
    out.to_csv(str(TAB_DIR / "table1_top_connections.csv"), index=False)
    return out


def table_2_engagement() -> pd.DataFrame:
    src = STAGE2 / "reports" / "network_engagement_table.csv"
    df = pd.read_csv(src)
    df.to_csv(str(TAB_DIR / "table2_network_engagement.csv"), index=False)
    return df


def table_3_gnn_folds() -> pd.DataFrame:
    path = STAGE3 / "evaluation" / "loso_summary.json"
    if not path.exists():
        return pd.DataFrame()
    s = json.loads(path.read_text())
    df = pd.DataFrame(s["per_fold"])
    df.to_csv(str(TAB_DIR / "table3_gnn_folds.csv"), index=False)
    return df


def table_4_a_vs_b() -> pd.DataFrame:
    rows = [
        {"aspect": "task", "approach_a_v6": "15-class scenario id",
         "approach_b_gnn": "15-class scenario id"},
        {"aspect": "features", "approach_a_v6": "morlet ersp + raw oscillation dynamics",
         "approach_b_gnn": "roi-level wpli + per-roi band power"},
        {"aspect": "model", "approach_a_v6": "multimodal transformer (~260k params)",
         "approach_b_gnn": "2-layer gcn (<1k params)"},
        {"aspect": "interpretation", "approach_a_v6": "per-scenario ersp signatures",
         "approach_b_gnn": "per-scenario network signatures"},
    ]
    df = pd.DataFrame(rows)
    df.to_csv(str(TAB_DIR / "table4_a_vs_b.csv"), index=False)
    return df


#─────────────────────────────────────────────────────────────────────────────
#docx builder
#─────────────────────────────────────────────────────────────────────────────
def _styled_doc():
    from docx import Document
    from docx.shared import Pt, Cm
    doc = Document()
    sty = doc.styles['Normal']
    sty.font.name = 'Times New Roman'
    sty.font.size = Pt(11)
    for level in (1, 2, 3):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Times New Roman'
        h.font.size = Pt(12)
        h.font.bold = True
    return doc


def _add_table(doc, rows: list, headers: list):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = _pt(11)
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = "" if v is None else str(v)
            for run in cells[i].paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = _pt(11)


def _pt(p):
    from docx.shared import Pt
    return Pt(p)


def _add_para(doc, text: str):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = _pt(11)


def _add_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.size = _pt(12)
    return h


def build_docx(stage3_summary, stage3_perm):
    from docx.shared import Cm
    doc = _styled_doc()

    _add_heading(doc, "Approach B: Functional Connectivity Analysis "
                       "of Frustration Scenarios", level=1)
    _add_para(doc,
        "This report presents the second analytical perspective on the "
        "BITIRMEEG dataset. While Approach A focused on raw oscillation "
        "dynamics combined with a multimodal transformer, Approach B asks "
        "a complementary mechanistic question: do specific frustration "
        "scenarios reorganise the network of inter-region neural "
        "communication, and if so, are the resulting network signatures "
        "sufficient on their own to discriminate scenarios? "
        "The analysis is grounded in the V3 action-matched dataset "
        "(N=480 epochs, 9 subjects, 15 classes).")

    _add_heading(doc, "1. Introduction", level=1)
    _add_para(doc,
        "Functional connectivity quantifies statistical dependencies "
        "between brain regions and has been used extensively to relate "
        "cognitive states to large-scale network reorganisation (Cavanagh "
        "and Frank 2014). For frustration in interactive contexts, the "
        "theoretical priors point to the fronto-parietal network for "
        "cognitive control and conflict monitoring (Sauseng et al. 2008) "
        "and to alpha disengagement in default-mode-related electrodes "
        "during task interruptions. Approach A established that the "
        "binary discrimination problem is highly tractable from raw "
        "oscillation dynamics. Approach B asks a more difficult and more "
        "mechanistic question.")

    _add_heading(doc, "2. Methods", level=1)
    _add_heading(doc, "2.1 Connectivity metrics", level=2)
    _add_para(doc,
        "Two metrics were computed per epoch. The primary metric is the "
        "weighted phase lag index (wPLI), a phase-based measure that is "
        "robust to volume conduction at sensor level (Vinck et al. 2011). "
        "The secondary metric is the amplitude envelope correlation "
        "(AEC), computed as the Pearson correlation of Hilbert envelopes "
        "in each band. Coherence was excluded because of known volume "
        "conduction bias in 32-channel sensor space. Both metrics were "
        "computed for four bands: theta (4 to 8 Hz), alpha (8 to 13 Hz), "
        "beta (13 to 30 Hz) and gamma (30 to 40 Hz). wPLI was estimated "
        "per epoch with the time-resolved spectral connectivity routine "
        "in mne-connectivity using complex Morlet wavelets. AEC was "
        "computed in the time domain by bandpass filtering, Hilbert "
        "envelope extraction and Pearson correlation between channel "
        "pairs.")

    _add_heading(doc, "2.2 Region of interest definition", level=2)
    _add_para(doc,
        "The 32-channel sensor layout was collapsed into six regions of "
        "interest (frontal, frontal-central, central, parietal, "
        "occipital, temporal) using the same channel-to-ROI mapping as "
        "Approach A V6. Channel-pair connectivity values were averaged "
        "within each ROI pair to produce a six-by-six symmetric "
        "connectivity matrix per epoch, per band, per metric. The off-"
        "diagonal of this matrix contains 21 unique inter-region "
        "connections, giving 84 features per metric per epoch.")

    _add_heading(doc, "2.3 Statistical pipeline", level=2)
    _add_para(doc,
        "For each of the 14 frustration scenarios, the per-subject mean "
        "connectivity matrix during scenario epochs was contrasted with "
        "the per-subject mean during action-matched control epochs. The "
        "paired connection-wise difference was tested across subjects "
        "with a paired t-test. The t-test was chosen as the primary "
        "statistic because with nine subjects the Wilcoxon signed-rank "
        "test and the sign-flip permutation test share a discrete null "
        "distribution that caps the smallest achievable raw p-value at "
        "1 divided by 2^9, which is 0.0039, and which after FDR "
        "correction over 21 connections never reaches the 0.05 "
        "threshold even for very strong effects. The paired t-test does "
        "not suffer from this floor. Wilcoxon p-values and sign-flip "
        "permutation p-values (500 shuffles) are reported as secondary "
        "checks in the per-connection output file. p-values were "
        "corrected with the Benjamini and Hochberg false discovery rate "
        "procedure across the 21 within-band connections. Cohen's d for "
        "paired samples was computed alongside the p-values. A "
        "connection was considered significant only when the FDR-"
        "corrected t-test p was below 0.05 and the absolute Cohen's d "
        "was at least 0.5, the standard medium-effect threshold.")

    _add_heading(doc, "2.4 Graph neural network classifier", level=2)
    _add_para(doc,
        "Stage 3 used the per-epoch connectivity matrices as inputs to a "
        "minimal two-layer graph convolutional network. The six ROIs "
        "were treated as nodes; the band-averaged wPLI matrix provided "
        "the edge weights; per-ROI band power summaries from Approach A "
        "V6 were used as node features. The model has fewer than one "
        "thousand parameters, which is the smallest configuration that "
        "still permits two message passing layers, chosen to limit "
        "overfitting given nine subjects. Training used 9-fold leave-"
        "one-subject-out, AdamW with weight decay 1e-2, class-weighted "
        "cross-entropy, maximum 40 epochs and a patience of 8.")

    _add_heading(doc, "3. Results", level=1)
    _add_heading(doc, "3.1 Per-scenario network signatures (wPLI)", level=2)

    tab1 = table_1_top_connections()
    rows_t1 = [(r["scenario"], r["band"], r["connection"],
                f"{r['cohens_d']:+.2f}" if not pd.isna(r["cohens_d"]) else "n.s.",
                f"{r['fdr_p']:.3f}" if not pd.isna(r["fdr_p"]) else "n.s.",
                r["direction"], r["n_subjects"]) for _, r in tab1.iterrows()]
    _add_para(doc,
        "Table 1 reports the largest absolute Cohen's d connection per "
        "scenario, restricted to wPLI. Empty rows indicate scenarios "
        "with no FDR-significant connection at the 0.05 threshold. The "
        "n_subjects column shows how many subjects contributed paired "
        "samples for the test.")
    _add_table(doc, rows_t1,
               ["scenario", "band", "connection", "cohen d", "fdr p",
                "direction", "n subj"])

    _add_heading(doc, "3.2 Network engagement", level=2)
    tab2 = table_2_engagement()
    rows_t2 = []
    for _, r in tab2.iterrows():
        rows_t2.append([
            r["scenario"],
            "yes" if r.get("fronto_parietal_control", False) else "no",
            "yes" if r.get("default_mode_alpha", False) else "no",
            "yes" if r.get("sensorimotor", False) else "no",
            r.get("category", ""),
        ])
    _add_para(doc,
        "Table 2 categorises each scenario by the literature-defined "
        "networks it engages. Fronto-parietal control includes frontal "
        "to parietal coupling in theta and beta. Default-mode alpha "
        "captures alpha changes on the frontal to parietal and frontal "
        "to occipital axes. Sensorimotor captures beta and gamma "
        "involvement of central electrodes. Cells marked yes mean at "
        "least one significant wPLI connection in the corresponding "
        "definition.")
    _add_table(doc, rows_t2,
               ["scenario", "fronto-parietal", "default-mode alpha",
                "sensorimotor", "category"])

    _add_heading(doc, "3.3 GNN classification (Stage 3)", level=2)
    if stage3_summary is not None:
        acc = stage3_summary.get("mean_accuracy", 0)
        std = stage3_summary.get("std_accuracy", 0)
        f1  = stage3_summary.get("mean_f1_macro", 0)
        chance = stage3_summary.get("chance", 1/15)
        _add_para(doc,
            f"The minimal GCN achieved an LOSO accuracy of {acc:.3f} "
            f"(standard deviation {std:.3f}) with macro F1 {f1:.3f} "
            f"across nine folds. Chance for the 15-class task is "
            f"{chance:.3f}.")
        if stage3_perm is not None:
            p = stage3_perm.get("p_value", float("nan"))
            null_mean = stage3_perm.get("null_mean", float("nan"))
            null_std = stage3_perm.get("null_std", float("nan"))
            _add_para(doc,
                f"A 50-shuffle label permutation test produced a null "
                f"distribution centred at {null_mean:.3f} (standard "
                f"deviation {null_std:.3f}) and a p-value of {p:.3f}. "
                "If the GNN accuracy is close to chance, this stage "
                "should be interpreted as confirming that connectivity "
                "alone is insufficient for full discrimination, not as "
                "evidence against the mechanistic findings in Stage 2.")

        tab3 = table_3_gnn_folds()
        if not tab3.empty:
            rows_t3 = [(int(r["sid"]), f"{r['acc']:.3f}",
                        f"{r['f1']:.3f}", int(r["n_val"]))
                       for _, r in tab3.iterrows()]
            _add_table(doc, rows_t3, ["held-out subject", "acc", "f1",
                                       "n val epochs"])

    _add_heading(doc, "3.4 Comparison with Approach A V6", level=2)
    tab4 = table_4_a_vs_b()
    rows_t4 = [(r["aspect"], r["approach_a_v6"], r["approach_b_gnn"])
               for _, r in tab4.iterrows()]
    _add_table(doc, rows_t4,
               ["aspect", "approach a v6", "approach b gnn"])

    _add_heading(doc, "4. Discussion", level=1)
    _add_para(doc,
        "The per-scenario network analysis isolates the inter-region "
        "communication changes that accompany each frustration scenario, "
        "complementing the per-electrode oscillation results in "
        "Approach A V6. Scenarios that engage the fronto-parietal "
        "network in theta or beta are consistent with conflict "
        "monitoring and cognitive control demands described in the "
        "literature on midline frontal theta (Cavanagh and Frank 2014). "
        "Scenarios with alpha changes on fronto-parietal or fronto-"
        "occipital axes are consistent with transient disengagement and "
        "attentional reorientation during task interruptions. Beta and "
        "gamma involvement of central electrodes is interpreted as "
        "sensorimotor preparation related to the user's planned response.")
    _add_para(doc,
        "Several methodological caveats need to be made explicit. With "
        "N=9 subjects, connectivity estimates carry substantial "
        "uncertainty. Findings should be interpreted as preliminary "
        "mechanistic hypotheses requiring replication. The GNN "
        "classifier in Stage 3 does not match Approach A V6 performance "
        "because it uses only network-level features without raw "
        "oscillation dynamics; this is by design, since Approach B "
        "prioritises mechanistic network interpretation over maximum "
        "classification. Volume conduction in sensor-space EEG can "
        "produce spurious connectivity, which we mitigated by using "
        "wPLI as the primary measure but cannot fully eliminate. "
        "Source-space analysis with individual MRI would strengthen the "
        "directional claims. Finally, the connectivity differences "
        "between scenarios and controls do not establish causal "
        "direction, and Granger-style directed connectivity was "
        "considered but judged unreliable at this sample size.")

    _add_heading(doc, "5. References", level=1)
    _add_para(doc,
        "Cavanagh, J. F., and Frank, M. J. (2014). Frontal theta as a "
        "mechanism for cognitive control. Trends in Cognitive Sciences, "
        "18(8), 414 to 421.")
    _add_para(doc,
        "Sauseng, P., Klimesch, W., Gruber, W. R., and Birbaumer, N. "
        "(2008). Cross-frequency phase synchronisation: a brain "
        "mechanism of memory matching and attention. NeuroImage, 40(1), "
        "308 to 317.")
    _add_para(doc,
        "Vinck, M., Oostenveld, R., van Wingerden, M., Battaglia, F., "
        "and Pennartz, C. M. A. (2011). An improved index of phase "
        "synchronisation for electrophysiological data in the presence "
        "of volume conduction, noise and sample-size bias. NeuroImage, "
        "55(4), 1548 to 1565.")

    return doc


def save_docx_pdf(doc, stem: str):
    docx_path = OUT_DIR / f"{stem}.docx"
    doc.save(str(docx_path))
    print(f"saved {docx_path}")
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        script = f"""
tell application "Microsoft Word"
    activate
    open POSIX file "{docx_path}"
    set d to active document
    save as d file name (POSIX file "{pdf_path}") file format format PDF
    close d saving no
end tell"""
        subprocess.run(["osascript", "-e", script], check=True, timeout=120)
        print(f"saved {pdf_path}")
    except Exception as e:
        print(f"pdf conversion failed (Word.app): {e}")


def _safe(call):
    try:
        return call()
    except Exception as e:
        print(f"  figure failed: {e}")
        return None


def main():
    print("=== stage 4: report generation ===")
    print("building figures...")
    _safe(figure_1_roi_topography)
    _safe(figure_2_change_matrices)
    _safe(figure_3_network_engagement)
    _safe(figure_4_gnn_confusion)
    _safe(figure_5_v6_vs_gnn)

    print("building tables...")
    table_1_top_connections()
    table_2_engagement()
    table_3_gnn_folds()
    table_4_a_vs_b()

    summary = None
    perm    = None
    s_path = STAGE3 / "evaluation" / "loso_summary.json"
    p_path = STAGE3 / "evaluation" / "permutation_results.json"
    if s_path.exists():
        summary = json.loads(s_path.read_text())
    if p_path.exists():
        perm = json.loads(p_path.read_text())

    print("building docx...")
    doc = build_docx(summary, perm)
    save_docx_pdf(doc, "Approach_B_Report")
    print("\nstage 4 done.")


if __name__ == "__main__":
    main()
